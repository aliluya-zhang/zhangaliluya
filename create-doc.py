# -*- coding: utf-8 -*-
"""生成《He Was Born》风格作品集网站的修改文档"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── 颜色常量 ──
C_BLUE_DARK  = RGBColor(0x1F, 0x4E, 0x79)
C_BLUE       = RGBColor(0x2E, 0x75, 0xB6)
C_BLUE_LIGHT = RGBColor(0xD5, 0xE8, 0xF0)
C_GRAY       = RGBColor(0x99, 0x99, 0x99)
C_LIGHT      = RGBColor(0xF5, 0xF8, 0xFC)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT       = RGBColor(0x33, 0x33, 0x33)

def hex_to_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC', size=4):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_table_border(table, color='CCCCCC', size=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1cm ≈ 567 twips
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

def cell_text(cell, text, bold=False, font_size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font_name='Arial'):
    """向单元格添加带格式的文本"""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run

def add_heading(doc, text, level=1, color=None):
    """添加带样式的标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = color or C_BLUE_DARK
        run.font.name = 'Arial'
        # 添加下划线
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E75B6')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = C_BLUE
        run.font.name = 'Arial'
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.name = 'Arial'
    return p

def add_para(doc, text, font_size=10, color=None, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, font_size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    return p

def add_number(doc, text, font_size=10):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Arial'
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(OxmlElement('w:pageBreakBefore'))
    return p

def make_header_row(table, headers, widths_cm, header_color='2E75B6'):
    """创建表头行"""
    header_row = table.rows[0]
    for i, (text, w) in enumerate(zip(headers, widths_cm)):
        cell = header_row.cells[i]
        set_cell_bg(cell, header_color)
        set_cell_borders(cell, header_color, 4)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_WHITE
        run.font.name = 'Arial'

def add_data_row(table, row_idx, cells_data, bg_even='FFFFFF', bg_odd='F5F8FC'):
    """添加数据行"""
    row = table.rows[row_idx + 1]  # +1 因为第一行是表头
    bg = bg_even if row_idx % 2 == 0 else bg_odd
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'CCCCCC', 4)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = C_TEXT

# ── 创建文档 ──
doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# 禁用默认样式干扰
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# ═══════════════════════════════════════════
# 封面 / 项目信息
# ═══════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('《He Was Born》风格作品集网站')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = C_BLUE_DARK
run.font.name = 'Arial'

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('前端开发修改文档')
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = C_BLUE
run2.font.name = 'Arial'

# 分隔线
pLine = doc.add_paragraph()
pLine.paragraph_format.space_after = Pt(16)
pPr = pLine._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '8')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '2E75B6')
pBdr.append(bottom)
pPr.append(pBdr)

# 元信息表格
meta_data = [
    ('文件名称', 'hewasborn-static.html'),
    ('文件路径', 'c:\\Users\\lili\\WorkBuddy\\20260509080323\\hewasborn-static.html'),
    ('开发日期', '2026-05-09'),
    ('版本', 'v1.0.0'),
    ('技术栈', 'HTML5 + CSS3 + JavaScript + Three.js CDN + Framer Motion CDN'),
]

meta_table = doc.add_table(rows=len(meta_data), cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (k, v) in enumerate(meta_data):
    bg = 'F2F7FC' if i % 2 == 0 else 'FFFFFF'
    cell0 = meta_table.rows[i].cells[0]
    cell1 = meta_table.rows[i].cells[1]
    set_cell_bg(cell0, bg)
    set_cell_bg(cell1, bg)
    set_cell_borders(cell0, 'DDDDDD', 2)
    set_cell_borders(cell1, 'DDDDDD', 2)
    r0 = cell0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.name = 'Arial'
    r0.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r1 = cell1.paragraphs[0].add_run(v)
    r1.font.size = Pt(10)
    r1.font.name = 'Arial'
    r1.font.color.rgb = C_TEXT

# 列宽
for row in meta_table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(14.0)

add_page_break(doc)

# ═══════════════════════════════════════════
# 一、项目概述
# ═══════════════════════════════════════════
add_heading(doc, '一、项目概述', level=1)
add_para(doc,
    '本项目复刻《He Was Born》网站（https://www.hewasborn.com/）的完整动效与交互设计，'
    '包含 9 大核心动画效果。项目采用「静态 HTML + CDN 引入」方式实现，无需编译环境，'
    '直接浏览器打开即可预览。设计风格为深色极简主题，主色调为柔和米金色 (#E8C4A0)，'
    '背景色 #0A0A0A，适合作品集展示。', font_size=10)

# ═══════════════════════════════════════════
# 二、技术选型
# ═══════════════════════════════════════════
add_heading(doc, '二、技术选型', level=1)

tech_data = [
    ('网页结构',     'HTML5 + CSS3 + Vanilla JavaScript（原生实现，无框架依赖）'),
    ('动画库',       'Framer Motion 11.15.0 (CDN) + Three.js 0.170.0 (CDN) + Canvas API'),
    ('字体',         'Inter (Google Fonts) + Noto Sans SC（中文回退）'),
    ('颜色主谱',     '背景 #0A0A0A / 文字 #FFFFFF / 主色 #E8C4A0 / 粉色 #FFB3D9 / 热情橙 #FF8A5C / 科技蓝 #6BB5FF'),
    ('响应式实现',   'CSS Media Queries（基于 768px 分断点）+ 移动端允许默认光标'),
]

tech_table = doc.add_table(rows=len(tech_data)+1, cols=2)
add_table_border(tech_table, 'CCCCCC')
headers = ['类别', '详情']
widths = [3.5, 14.0]
make_header_row(tech_table, headers, widths, '2E75B6')
for i, (k, v) in enumerate(tech_data):
    add_data_row(tech_table, i, [k, v])
# 列宽
for row in tech_table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(14.0)

add_page_break(doc)

# ═══════════════════════════════════════════
# 三、全部动效清单
# ═══════════════════════════════════════════
add_heading(doc, '三、全部动效清单（共 9 项）', level=1)

effects_data = [
    ['#1', 'Logo 波光动画',   '#wave-canvas (Hero)', '页面加载时 Canvas 绘制 3 层正弦波带，blur 滤镜叠加，透明度从 0.02 渐变到 0.06，营造水面反射柔和光感。持续循环运行。'],
    ['#2', '自定义光标跟随',  '#cursor',             '监听 mousemove 事件，绘制 20px 白圆形，mix-blend-mode: difference。hover 可点击元素时放大至 44px 并变色。移动端隐藏。'],
    ['#3', '项目卡片悬停',    '.project-card',        'CSS transition: scale(1.03)，图片 scale(1.08)，边框高亮，box-shadow 阴影。持续 0.25s ease-out。'],
    ['#4', '汉堡菜单滑入',    '#menu-panel',         'translateX(100%) 滑入，背景浅黑渐透。菜单项通过 nth-child 延迟递进（0.05s/项）。'],
    ['#5', 'Three.js 背景',  '#bg-canvas',           'WebGL 渲染 Icosahedron / Octahedron / Tetrahedron / Torus 线框几何体，粉/橙/蓝高饱和色，屏幕固定在背景，持续旋转。'],
    ['#6', '横向滚动视差',   '#works-track',         '监听 scroll 事件，将滚动进度映射到 translateX，使卡片列横向移动，产生视差效果。需手动滚动才能观察。'],
    ['#7', '滚动触发动画',   '.reveal',              'IntersectionObserver 监听元素进入视口（threshold 0.15），添加 in-view 类后触发 CSS 转换：opacity 0→1、translateY 50px→0，交错延迟 0.1s。'],
    ['#8', '数字递增计数器', '.counter-num',         'IntersectionObserver 触发后用 requestAnimationFrame 运行 easeOutQuart 缓动，1.5s 内从 0 递增至目标值，每个数字独立运行。'],
    ['#9', '橡皮筋回弹效果', '.elastic-wrap',        '滚动到顶部/底部时继续滚动产生弹性位移（translateY），超出量为实际超出的 30%，使用 scroll 事件被动监听实现。'],
]

eff_table = doc.add_table(rows=len(effects_data)+1, cols=4)
add_table_border(eff_table, 'CCCCCC')
make_header_row(eff_table, ['编号', '动效名称', '位置', '实现详情'], [1.2, 3.0, 3.5, 9.8])
for i, row_data in enumerate(effects_data):
    add_data_row(eff_table, i, row_data)
for row in eff_table.rows:
    row.cells[0].width = Cm(1.2)
    row.cells[1].width = Cm(3.0)
    row.cells[2].width = Cm(3.5)
    row.cells[3].width = Cm(9.8)

add_page_break(doc)

# ═══════════════════════════════════════════
# 四、页面结构
# ═══════════════════════════════════════════
add_heading(doc, '四、页面结构', level=1)

structure_data = [
    ['导航栏',      '#navbar',              '导航栏，滚动 60px 后添加毛玻璃背景（backdrop-filter blur）'],
    ['汉堡菜单',    '#menu-overlay + #menu-panel', '滑出菜单面板，从右侧滑入，配合背景遮罩'],
    ['Hero 首屏',  '#hero',                 '波光 Canvas 背景，文字 2.2s 延迟后淡入，底部居中「Hi~ I\'m 阿莉路亚」'],
    ['关于我们',    '#about',                '双列布局，右侧点阵装饰图，滚动淡入动画'],
    ['服务专页',    '#services',             '三栏卡片布局，悬停浮起效果'],
    ['作品展示',    '#works',                '横向滚动区，项目卡片网格'],
    ['数据计数器',  '.counter-row',         '三个统计数字，滚动触发的递增动画'],
    ['联系我们',    '#contact',             '居中 CTA 按钮，邮箱链接'],
    ['页脚',        '.footer',              '版权信息和技能标签'],
]

struct_table = doc.add_table(rows=len(structure_data)+1, cols=3)
add_table_border(struct_table, 'CCCCCC')
make_header_row(struct_table, ['区块', 'ID', '内容说明'], [3.0, 4.5, 10.0])
for i, row_data in enumerate(structure_data):
    add_data_row(struct_table, i, row_data)
for row in struct_table.rows:
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(10.0)

add_page_break(doc)

# ═══════════════════════════════════════════
# 五、修改指南
# ═══════════════════════════════════════════
add_heading(doc, '五、修改指南', level=1)

add_heading(doc, '5.1 如何修改文字内容', level=2)
add_bullet(doc, '在 HTML 文件中查找相应的文字内容，直接修改即可。')
add_bullet(doc, '项目卡片内容位于 works-track 区块中，每个 project-card 包含标题、描述和标签。')
add_bullet(doc, '字体描述位于 .hero-sub 和 .hero-title 元素中。')

add_heading(doc, '5.2 如何替换背景视差', level=2)
add_para(doc, '当前采用 Three.js 绘制动态几何体背景。如需替换为固定图片，修改 CSS 中的 #bg-canvas 相关代码：', font_size=10)
p_code = doc.add_paragraph()
p_code.paragraph_format.left_indent = Cm(0.75)
p_code.paragraph_format.space_after = Pt(4)
run_code = p_code.add_run('#bg-canvas { background: url(\'your-image.jpg\') center/cover no-repeat fixed; }')
run_code.font.name = 'Consolas'
run_code.font.size = Pt(8)
run_code.font.color.rgb = C_BLUE

add_heading(doc, '5.3 如何添加新的项目卡片', level=2)
add_number(doc, '在 works-track 元素内添加新的 project-card div')
add_number(doc, '修改 project-thumb-inner 的 background 属性来调整颜色')
add_number(doc, '添加 data-cursor-hover 属性以启用自定义光标效果')

add_heading(doc, '5.4 颜色主谱修改', level=2)
add_para(doc, '在 <style> 区块中查找 .text-gradient，修改 background 属性即可调整渐变文字颜色。颜色变量存储在 CSS :root 中，可添加 --primary-color 等自定义属性。', font_size=10)

# ═══════════════════════════════════════════
# 六、已知问题
# ═══════════════════════════════════════════
add_heading(doc, '六、已知限制与问题', level=1)
add_bullet(doc, '移动端隐藏自定义光标，使用默认光标（@media (pointer: coarse)）')
add_bullet(doc, '横向滚动需手动滚动才能观察，滚轮无法触发')
add_bullet(doc, 'Three.js 背景在极低速设备上可能有性能问题，建议移动端使用静态背景')
add_bullet(doc, '此项目为预览版本，正式生产环境建议迁移至 Next.js App Router 项目')

# ═══════════════════════════════════════════
# 七、技术架构说明
# ═══════════════════════════════════════════
add_heading(doc, '七、技术架构说明', level=1)
add_heading(doc, '网页层次结构', level=2)
p_layer = doc.add_paragraph()
p_layer.paragraph_format.left_indent = Cm(0.75)
p_layer.paragraph_format.space_after = Pt(8)
run_layer = p_layer.add_run('body > #cursor (z:9999) > #bg-canvas (z:0 fixed) > .elastic-wrap > nav + sections + footer')
run_layer.font.name = 'Consolas'
run_layer.font.size = Pt(8)
run_layer.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_para(doc, '层次顺序：光标最高（控制元素），海树等图层（视觉效果），内容层（可滚动），固定背景（最低）', font_size=10)

add_heading(doc, 'JavaScript 函数分栏', level=2)
func_data = [
    ('initCursor()',        '自定义光标跟随动画 + hover 状态切换'),
    ('initWave()',          'Canvas 绘制 Hero 区域波光动画'),
    ('initMenu()',          '汉堡菜单打开/关闭逻辑'),
    ('initThreeBG()',       'Three.js WebGL 渲染器和线框几何体旋转'),
    ('initHorizontalScroll()', '横向港动联动，滚动进度转换'),
    ('initReveal()',        'IntersectionObserver 滚动介入触发'),
    ('initCounters()',      'requestAnimationFrame 数字递增计数器'),
    ('initElastic()',       '橡皮筋效果，滚动超界时的弹性偏移'),
]

func_table = doc.add_table(rows=len(func_data)+1, cols=2)
add_table_border(func_table, 'CCCCCC')
make_header_row(func_table, ['函数名称', '职责'], [4.5, 13.0])
for i, row_data in enumerate(func_data):
    row = func_table.rows[i+1]
    bg = 'FFFFFF' if i % 2 == 0 else 'F5F8FC'
    for j, text in enumerate(row_data):
        cell = row.cells[j]
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'CCCCCC', 4)
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        if j == 0:
            run.bold = True
            run.font.color.rgb = C_BLUE
        else:
            run.font.color.rgb = C_TEXT
for row in func_table.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(13.0)

add_page_break(doc)

# ═══════════════════════════════════════════
# 八、后续开发计划
# ═══════════════════════════════════════════
add_heading(doc, '八、后续开发计划', level=1)
add_number(doc, '安装 Node.js 后，将项目迁移至 Next.js 15 App Router + TypeScript + Tailwind CSS + Framer Motion 正式项目结构')
add_number(doc, '添加页面转场转转的入场动画（基于 Framer Motion LayoutId）')
add_number(doc, '完善响应式实现，正式支持移动端屏幕尺寸自适应')
add_number(doc, '添加更多项目存储在 projects.ts 配置文件中，实现数据与表现分离')

# 结尾
p_end = doc.add_paragraph()
p_end.paragraph_format.space_before = Pt(30)
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run('— 文档结束 —')
run_end.italic = True
run_end.font.size = Pt(9)
run_end.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run_end.font.name = 'Arial'

# ── 保存 ──
output_path = r'c:\Users\lili\WorkBuddy\20260509080323\hewasborn-modification-report.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
