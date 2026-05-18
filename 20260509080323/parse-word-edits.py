# -*- coding: utf-8 -*-
"""Parse Word document and output all edits (col2 ≠ col3)"""
from docx import Document
try:
    from docx.oxml.ns import qn
except ImportError:
    from docx.oxml import qn
import sys, json, re

DOCX = r"C:\Users\lili\WorkBuddy\20260509080323\aliluya-content-edit.docx"

def cell_text(cell):
    return ''.join(p.text for p in cell.paragraphs).strip()

doc = Document(DOCX)
edits = []

for t_idx, table in enumerate(doc.tables):
    rows = table.rows
    for r_idx, row in enumerate(rows):
        cells = row.cells
        # 2-column table: col0=current, col1=edit
        if len(cells) >= 2:
            cur = cell_text(cells[0])
            new = cell_text(cells[1]) if len(cells) > 1 else ''
            if new and new != cur:
                edits.append({
                    'table': t_idx,
                    'row':   r_idx,
                    'current': cur,
                    'new':   new,
                })
        # 4-column table (masonry): col2=title edit, col3=desc edit
        if len(cells) >= 4:
            for c_idx in [2, 3]:
                cur = cell_text(cells[c_idx])
                # "内容区（可编辑）" is not real content
                if cur and cur not in ('内容区（可编辑）', '图片路径（绝对路径）', '大标题（hover 显示）', '小标题（hover 显示）'):
                    edits.append({
                        'table': t_idx,
                        'row':   r_idx,
                        'col':   c_idx,
                        'current': '-',
                        'new':   cur,
                    })

if edits:
    print(f"找到 {len(edits)} 处修改：")
    for e in edits:
        print(f"  Table{e['table']} Row{e['row']} | 新值：{e['new'][:60]}")
else:
    print("未检测到修改（第三列均为空或与当前值相同）")
    print("提示：请确认 Word 文件已保存，且第三列确实有内容。")
