# -*- coding: utf-8 -*-
"""Apply Word edits to HTML: for each table, compare col0 vs col1 (2-col),
   or col2/col3 (4-col masonry table).  Output a JSON patch file."""
from docx import Document
import json, re, sys

DOCX = r"C:\Users\lili\WorkBuddy\20260509080323\aliluya-content-edit.docx"
HTML = r"c:\Users\lili\WorkBuddy\20260509080323\hewasborn-static.html"
PATCH= r"c:\Users\lili\WorkBuddy\20260509080323\word-patch.json"

SKIP = {
    '内容区（可编辑）', '图片路径（绝对路径）',
    '大标题（hover 显示）', '小标题（hover 显示）',
    '本文档路径：', 'C:\\Users\\lili\\WorkBuddy\\20260509080323',
}

def cell_text(cell):
    return ''.join(p.text for p in cell.paragraphs).strip()

def normalize(s):
    return re.sub(r'\s+', ' ', s.strip())

doc  = Document(DOCX)
with open(HTML, encoding='utf-8') as f:
    html = f.read()

patches = []
for ti, tbl in enumerate(doc.tables):
    rows = tbl.rows
    nc   = len(rows[0].cells)
    for ri, row in enumerate(rows):
        cells = row.cells
        # 2-column table: col1 = new value
        if nc == 2:
            cur = cell_text(cells[0])
            new = cell_text(cells[1])
            if new and new not in SKIP and normalize(new) != normalize(cur):
                patches.append({'table':ti,'row':ri,'type':'2col',
                               'current':cur,'new':new})
        # 4-column table: col2=title edit, col3=desc edit
        if nc == 4:
            for ci in [2,3]:
                new = cell_text(cells[ci])
                if new and new not in SKIP:
                    patches.append({'table':ti,'row':ri,'col':ci,'type':'4col',
                                   'new':new})

# Try to apply patches to HTML
applied = []
for p in patches:
    new = p['new']
    # Skip if already in HTML
    if normalize(new) in normalize(html):
        applied.append({'...': 'skip-exists', 'new': new[:40]})
        continue

    if p['type'] == '2col':
        cur = p['current']
        if cur in html:
            html = html.replace(cur, new, 1)
            applied.append({'ok':'2col', 'row':p['row'], 'new':new[:50]})
        else:
            applied.append({'MISS': '2col', 'cur':cur[:40], 'new':new[:40]})
    elif p['type'] == '4col':
        # For masonry: find nearby context and insert
        applied.append({'4col': p['row'], 'col': p['col'], 'new': new[:40]})

with open(HTML, 'w', encoding='utf-8') as f:
    f.write(html)
with open(PATCH, 'w', encoding='utf-8') as f:
    json.dump({'patches': patches, 'applied': applied}, f, ensure_ascii=False, indent=2)

print(f"Patches: {len(patches)}, Applied: {len([a for a in applied if 'ok' in a])}")
print("Details →", PATCH)
