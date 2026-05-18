# -*- coding: utf-8 -*-
"""根据 hewasborn-static.html 当前内容重新生成 aliluya 内容修改文档
   第 3 列（内容区；可编辑）留空白，带虚线边框
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\lili\WorkBuddy\20260509080323\aliluya-content-edit.docx"
doc = Document()

# ── 页面边距 ──
s = doc.sections[0]
s.top_margin = s.bottom_margin = Cm(2.5)
s.left_margin = s.right_margin = Cm(2.5)

# ── 工具函数 ────────────────────────────────────────────
def set_bg(cell, hex6):
    tc = cell._tc
    p  = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    p.append(shd)

def cell_write(cell, text, bold=False, sz=9.5, color='1a1a1a',
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name      = 'Arial'
    run.font.size      = Pt(sz)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rF   = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), 'Arial')
    rPr.insert(0, rF)

def add_line(text='', bold=False, sz=10, color='1a1a1a', italic=False,
             sb=6, sa=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run = p.add_run(text)
        run.font.name      = 'Arial'
        run.font.size      = Pt(sz)
        run.font.bold      = bold
        run.font.italic    = italic
        run.font.color.rgb = RGBColor.from_string(color)
        rPr = run._r.get_or_add_rPr()
        rF   = OxmlElement('w:rFonts')
        rF.set(qn('w:eastAsia'), 'Arial')
        rPr.insert(0, rF)
    return p

def divider():
    add_line('─' * 90, sz=7, color='CCCCCC', sb=4, sa=4)

def section_tag(text):
    add_line(text.upper(), bold=True, sz=8.5, color='AAAAAA', sb=18, sa=4)

def h1(text):
    add_line(text, bold=True, sz=11, color='1a1a1a', sb=8, sa=4)

def hint(text):
    add_line(text, italic=True, sz=8, color='999999', sb=2, sa=4)

def spacer():
    add_line('', sb=6, sa=6)

def add_dashed_border(cell, color='CCCC00'):
    """给单元格加虚线边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'dashed')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcPr.append(b)

# ── 核心：生成 3 列表格（元素名 | 当前内容 | 内容区·可编辑） ──
def edit_table(rows_data, col_w=None, header=None):
    """
    rows_data: [(label, current_value), ...]
    生成 3 列：0=元素名  1=当前内容  2=内容区（空白可编辑）
    """
    n_cols = 3
    if col_w is None:
        col_w = [Cm(3.8), Cm(7.5), Cm(5.7)]
    t = doc.add_table(rows=len(rows_data) + (1 if header else 0), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ci, w in enumerate(col_w):
        for cell in t.columns[ci].cells:
            cell.width = w

    row_offset = 0
    if header:
        for ci, h in enumerate(header):
            c = t.rows[0].cells[ci]
            set_bg(c, '1a1a1a')
            cell_write(c, h, bold=True, sz=8, color='FFFFFF',
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        row_offset = 1

    for r_idx, (label, cur_val) in enumerate(rows_data):
        actual_r = r_idx + row_offset
        bg = 'F5F5F5' if r_idx % 2 == 0 else 'FAFAFA'
        row_cells = t.rows[actual_r].cells

        # 第 0 列：元素名
        set_bg(row_cells[0], bg)
        cell_write(row_cells[0], label, bold=False, sz=9, color='1a1a1a')
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # 第 1 列：当前内容
        set_bg(row_cells[1], 'FFFFFF')
        cell_write(row_cells[1], cur_val, sz=8.5, color='1a1a1a')
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # 第 2 列：内容区（空白可编辑）── 淡黄底色 + 虚线边框
        set_bg(row_cells[2], 'FFFFF0')
        # 不留任何文字，只加虚线边框提示
        add_dashed_border(row_cells[2], 'CCCC00')
        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    spacer()
    return t

# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
add_line('ALILUYA', bold=True, sz=38, color='1a1a1a', sb=30, sa=2)
add_line('作品集网站  ·  内容修改文档', sz=13, color='666666', sb=0, sa=2)
add_line('Portfolio Website  ·  Content Edit Guide', sz=9, color='AAAAAA',
         italic=True, sb=0, sa=16)
divider()

# 文件路径表（2 列，不调用 edit_table）
t_fp = doc.add_table(rows=4, cols=2)
t_fp.alignment = WD_TABLE_ALIGNMENT.LEFT
col_w_fp = [Cm(4.2), Cm(12.8)]
for i, w in enumerate(col_w_fp):
    for cell in t_fp.columns[i].cells:
        cell.width = w

fp_rows = [
    ('📁 HTML 文件路径',   r'C:\Users\lili\Desktop\hewasborn-static.html'),
    ('🎬 背景视频路径',    r'C:\Users\lili\Desktop\豆包.mp4'),
    ('🖼️ 瀑布流图片目录',  r'C:\Users\lili\WorkBuddy\20260509080323\img\  （含 1.jpg ~ 6.jpg）'),
    ('📝 本文档路径',      r'C:\Users\lili\Desktop\aliluya-content-edit.docx'),
]
for i, (label, val) in enumerate(fp_rows):
    bg = 'F2F2F2' if i % 2 == 0 else 'F8F8F8'
    c0 = t_fp.rows[i].cells[0]
    c1 = t_fp.rows[i].cells[1]
    set_bg(c0, bg)
    set_bg(c1, 'FFFFFF')
    cell_write(c0, label, bold=True, sz=9, color='555555')
    cell_write(c1, val, sz=8.5, color='1a1a1a')
spacer()
divider()
add_line('💡 编辑方法：在第 3 列【内容区（可编辑）】的空白格中直接输入文字，',
         italic=True, sz=8.5, color='888888', sb=4, sa=2)
add_line('   图片路径填完整绝对路径。颜色格式：HEX（#e8e9eb）或 rgba（rgba(26,26,26,0.75)）。',
         italic=True, sz=8.5, color='888888', sb=0, sa=4)
add_line('   修改完成后保存本文档，发给我，我会同步到 HTML。',
         italic=True, sz=8.5, color='888888', sb=0, sa=20)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 1 部分：导航栏
# ═══════════════════════════════════════════════════════════
section_tag('第 1 部分')
add_line('导航栏  ·  Navbar', bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 325-344 行（桌面导航） · 第 348-358 行（汉堡菜单）')
spacer()

h1('1.1  Logo 文字')
edit_table([('Logo 文字', 'Aliluya')])
hint('HTML 第 326 行：<div class="nav-logo">Aliluya</div>  ·  桌面端左上角 Logo')
spacer()

h1('1.2  导航下拉触发文字 + 6 个下拉项')
edit_table([
    ('下拉触发文字',   '项目案例'),
    ('下拉项 ①',      '品牌设计'),
    ('下拉项 ②',      '网页开发'),
    ('下拉项 ③',      '3D 视觉'),
    ('下拉项 ④',      '视频制作'),
    ('下拉项 ⑤',      'UI/UX 设计'),
    ('下拉项 ⑥',      '动态交互'),
])
hint('HTML 第 329 行（桌面下拉触发） · 第 350-355 行（桌面下拉内容） · 第 360-365 行（汉堡菜单）需同步修改')
spacer()

h1('1.3  导航联系文字')
edit_table([('导航联系文字', '联系我')])
hint('HTML 第 339 行：<a href="#contact">联系我</a>  ·  同时更新汉堡菜单第 366 行')
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 2 部分：Hero 首屏
# ═══════════════════════════════════════════════════════════
section_tag('第 2 部分')
add_line('Hero 首屏  ·  Hero Section', bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 359-377 行  ·  背景视频：与 HTML 同目录')
spacer()

h1('2.1  主标题')
edit_table([('当前内容', 'Hi~\nI\'m 阿莉路亚')])
hint('HTML 第 367 行：<span class="hero-name">阿莉路亚</span>  ·  <br> 为换行符')
spacer()

h1('2.2  正文第一段')
edit_table([('当前内容',
    '8年以上设计行业经验，横跨数字艺术、游戏视觉及产品全案设计。'
    '曾参与王者好物和王者荣耀造物节七周年KV海报制作，'
    '自创 IP 曾被上市公司独家买断一年。')])
hint('HTML 第 369 行第一段（<br><br> 之间的内容）')
spacer()

h1('2.3  正文第二段')
edit_table([('当前内容',
    '我擅长将品牌逻辑与艺术感知结合，曾主导美容仪器产品的配色设计、'
    '包装配套及全链路物料落地。从 0 到 1 构建品牌视觉体系，'
    '不仅能画好一张画，更能打磨好一个产品。')])
hint('HTML 第 370 行第二段（<br><br> 之后的内容）')
spacer()

h1('2.4  背景视频路径')
edit_table([
    ('当前路径',   r'C:\Users\lili\Desktop\豆包.mp4'),
    ('推荐格式',   'MP4，尺寸 1920×1080 或更高，文件小于 50MB'),
])
hint('HTML 第 363 行：<source src="...">  ·  建议与 HTML 文件放在同一文件夹')
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 3 部分：About
# ═══════════════════════════════════════════════════════════
section_tag('第 3 部分')
add_line('About 关于区  ·  About Section', bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 379-403 行')
spacer()

h1('3.1  小标签（section-label）')
edit_table([('当前内容', 'About')])
hint('HTML 第 382 行：<p class="section-label">About</p>')
spacer()

h1('3.2  大标题（section-title）')
edit_table([('当前内容', 'Crafting Digital\nExperiences')])
hint('HTML 第 383 行：<h2 class="section-title">  ·  用 <br> 换行')
spacer()

h1('3.3  正文第一段')
edit_table([('当前内容',
    '我是一名专注于创意前端开发与设计的技术美术。'
    '擅长将先进技术（WebGL、Framer Motion、Three.js）与审美直觉结合，'
    '打造令人难忘的数字体验。曾参与《王者荣耀》海报设计，创作的 IP 被上市公司收购。')])
hint('HTML 第 387-392 行：<p>...第一段内容...</p>')
spacer()

h1('3.4  正文第二段')
edit_table([('当前内容',
    '目前主导 NenoVia/霓诺（FDA/CE 认证医美品牌）全案设计，'
    '同时推进多个创意作品集项目。')])
hint('HTML 第 394-397 行：<p>...第二段内容...</p>')
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 4 部分：瀑布流（4 列表格：# | 图片路径 | 大标题 | 小标题）
# ═══════════════════════════════════════════════════════════
section_tag('第 4 部分')
add_line('NenoVia 瀑布流图片区  ·  Masonry Gallery',
         bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 405-451 行  ·  布局：CSS columns，3 列，高度由图片原始比例决定')
spacer()

h1('4.1  6 张瀑布流卡片（图片路径 + 文字）')
# 4 列表格：# | 图片路径（当前值）| 大标题（可编辑）| 小标题（可编辑）
t_m = doc.add_table(rows=7, cols=4)
t_m.alignment = WD_TABLE_ALIGNMENT.LEFT
m_headers = ['#', '图片路径（当前值）', '大标题（hover 显示）\n⬜ 可编辑',
             '小标题（hover 显示）\n⬜ 可编辑']
m_col_w = [Cm(0.7), Cm(5.5), Cm(4.5), Cm(5.8)]
for ci, (h, w) in enumerate(zip(m_headers, m_col_w)):
    c = t_m.rows[0].cells[ci]
    c.width = w
    set_bg(c, '1a1a1a')
    cell_write(c, h, bold=True, sz=8, color='FFFFFF',
               align=WD_ALIGN_PARAGRAPH.CENTER)

m_rows = [
    ('1', r'C:\Users\lili\WorkBuddy\20260509080323\img\1.jpg'),
    ('2', r'C:\Users\lili\WorkBuddy\20260509080323\img\3.jpg'),
    ('3', r'C:\Users\lili\WorkBuddy\20260509080323\img\4.jpg'),
    ('4', r'C:\Users\lili\WorkBuddy\20260509080323\img\5.jpg'),
    ('5', r'C:\Users\lili\WorkBuddy\20260509080323\img\6.jpg'),
    ('6', r'C:\Users\lili\WorkBuddy\20260509080323\img\1.jpg（复用）'),
]
for r_idx, (num, path) in enumerate(m_rows):
    bg = 'F5F5F5' if r_idx % 2 == 0 else 'FAFAFA'
    row_cells = t_m.rows[r_idx + 1].cells

    # 列0: # 编号
    set_bg(row_cells[0], bg)
    cell_write(row_cells[0], num, sz=8.5, color='555555',
               align=WD_ALIGN_PARAGRAPH.CENTER)
    row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 列1: 图片路径（当前值，只读，白色背景）
    set_bg(row_cells[1], 'FFFFFF')
    cell_write(row_cells[1], path, sz=8, color='1a1a1a')
    row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 列2: 大标题（可编辑——留空 + 虚线边框 + 淡黄底色）
    set_bg(row_cells[2], 'FFFFF0')
    # 故意不写文字，留空
    add_dashed_border(row_cells[2], 'CCCC00')
    row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # 列3: 小标题（可编辑——留空 + 虚线边框 + 淡黄底色）
    set_bg(row_cells[3], 'FFFFF0')
    # 故意不写文字，留空
    add_dashed_border(row_cells[3], 'CCCC00')
    row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.TOP

hint('注：第 6 张目前复用 1.jpg，若有 2.jpg 可替换。大标题/小标题列带有 🟡 虚线边框，可直接输入新内容。')
spacer()

h1('4.2  瀑布流 CSS 样式参数（只读）')
edit_table([
    ('瀑布流列数',    'columns: 3 300px（3 列，最小宽度 300px）'),
    ('卡片圆角',      'border-radius: 16px'),
    ('卡片间距',      'margin-bottom: 16px'),
    ('hover 遮罩背景', '从透明 → 半透明 rgba(0,0,0,0.75)'),
    ('hover 文字颜色', '白色 #ffffff'),
])
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 5 部分：Contact
# ═══════════════════════════════════════════════════════════
section_tag('第 5 部分')
add_line('Contact 联系区  ·  Contact Section',
         bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 457-467 行')
spacer()

edit_table([
    ('小标签（section-label）',  'Contact'),
    ('大标题（section-title）',  "Let's Work\nTogether"),
    ('描述文字（contact-sub）',  '有想法？一起把数字体验做到极致。'),
    ('按钮文字',                 '发邮件给我 →'),
    ('邮箱地址（mailto）',       'mailto:ariluya@qq.com'),
])
hint('大标题用 <br> 换行；邮箱支持 mailto: 或 https:// 格式')
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 6 部分：Footer
# ═══════════════════════════════════════════════════════════
section_tag('第 6 部分')
add_line('Footer 页脚  ·  Footer', bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 469-473 行')
spacer()

edit_table([
    ('版权文字',   '© 2026 阿莉路亚. All rights reserved.'),
    ('技术栈说明', 'Built with Next.js · Three.js · Framer Motion'),
])
spacer()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 附录：全局样式
# ═══════════════════════════════════════════════════════════
section_tag('附录')
add_line('全局样式参数  ·  Global Styles',
         bold=True, sz=12, color='1a1a1a', sb=4, sa=2)
hint('HTML 第 1-330 行（<style> 区块）')
spacer()

edit_table([
    ('页面背景色',          '#e8e9eb（浅灰）'),
    ('正文文字色',          '#1a1a1a'),
    ('小标签颜色',          'rgba(26,26,26,0.35)'),
    ('光标颜色',            'rgba(26,26,26,0.75)'),
    ('光标尺寸',            '20px（hover 时 40px）'),
    ('Three.js 几何透明度', 'opacity: 0.06'),
    ('波光 Canvas 透明度',  'rgba(150,150,150, 0.015)'),
    ('字体（标题）',        'Kanit'),
    ('字体（正文）',        'Plus Jakarta Sans'),
    ('图片高度自适应',       'height: auto（由原始比例决定）'),
])

spacer()
divider()
add_line('', sb=4, sa=2)
add_line('✏️ 修改方法：在 🟡 虚线边框的【内容区】格子里直接输入新内容，保存后发给我。',
         italic=True, sz=9, color='888888', sb=0, sa=4)
add_line('   每次可修改多个区域，一起告知我即可，不需要逐个提交。',
         italic=True, sz=9, color='888888', sb=0, sa=10)

doc.save(OUTPUT)
print(f'OK -> {OUTPUT}')
